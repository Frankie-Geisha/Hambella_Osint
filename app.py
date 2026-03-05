import streamlit as st
import requests
from bs4 import BeautifulSoup
from openai import OpenAI
import anthropic 
from docx import Document
from io import BytesIO
import json
import os
from duckduckgo_search import DDGS  # 🌟 破壁行动：开源信息核查引擎
from supabase import create_client, Client

# ==========================================
# 🌸 1. 网页基础与【页面路由及身份系统】
# ==========================================
st.set_page_config(page_title="花魁 OSINT", page_icon="🌸", layout="wide", initial_sidebar_state="expanded")

# ==========================================
# 🌸 V6.0 商业级 UI 视觉覆写 (修复侧边栏 + Widget 级内饰排版)
# ==========================================
st.markdown("""
<style>
    /* 1. 全局底色与冗余清理 (精准隐蔽，保留侧边栏控制权！) */
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    .stDeployButton {display: none;} /* 只隐藏 Deploy 按钮，不隐藏整个 Header */
    [data-testid="stHeader"] {background: transparent !important;} /* 头部透明化 */
    .block-container {padding-top: 2rem; padding-bottom: 2rem;}

    .stApp {
        background-color: #F8FAFC !important; /* 极浅的蓝灰色背景 */
        background-image: radial-gradient(at 0% 0%, hsla(210, 100%, 98%, 1) 0px, transparent 50%),
                          radial-gradient(at 100% 0%, hsla(340, 100%, 98%, 1) 0px, transparent 50%);
    }

    /* 🌟 修复并升级侧边栏：磨砂亚克力质感 🌟 */
    [data-testid="stSidebar"] {
        background-color: rgba(248, 250, 252, 0.75) !important;
        backdrop-filter: blur(24px) !important;
        -webkit-backdrop-filter: blur(24px) !important;
        border-right: 1px solid rgba(0, 0, 0, 0.05) !important;
    }

    /* 2. 核心情报卡片外壳 (Widget 化) */
    [data-testid="stVerticalBlockBorderWrapper"] {
        background: rgba(255, 255, 255, 0.75) !important; 
        backdrop-filter: blur(20px) !important;
        -webkit-backdrop-filter: blur(20px) !important;
        border: 1px solid rgba(255, 255, 255, 0.8) !important;
        border-radius: 20px !important;
        box-shadow: 0 4px 15px rgba(0, 0, 0, 0.03), inset 0 1px 0 rgba(255, 255, 255, 1) !important;
        transition: transform 0.4s cubic-bezier(0.34, 1.56, 0.64, 1), box-shadow 0.4s ease !important;
        padding: 0.5rem !important; 
    }
    
    [data-testid="stVerticalBlockBorderWrapper"]:hover {
        transform: translateY(-4px) scale(1.005) !important;
        box-shadow: 0 12px 24px rgba(0, 0, 0, 0.06), inset 0 1px 0 rgba(255, 255, 255, 1) !important;
    }

    /* 3. 卡片内饰重塑 (文字、行距、层级) */
    
    [data-testid="stVerticalBlockBorderWrapper"] h2, 
    [data-testid="stVerticalBlockBorderWrapper"] h3, 
    [data-testid="stVerticalBlockBorderWrapper"] h4 {
        color: #1E293B !important;
        font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, Helvetica, Arial, sans-serif !important;
        font-weight: 700 !important;
        letter-spacing: -0.5px !important; 
        margin-bottom: 0.2rem !important;
    }

    [data-testid="stVerticalBlockBorderWrapper"] [data-testid="stMarkdownContainer"] p {
        color: #475569 !important; 
        font-size: 15px !important;
        line-height: 1.65 !important; 
    }

    [data-testid="stCaptionContainer"] {
        color: #94A3B8 !important;
        font-size: 13px !important;
        font-weight: 500 !important;
        margin-bottom: 0.5rem !important;
    }

    hr {
        border: none !important;
        height: 1px !important;
        background: linear-gradient(to right, rgba(0,0,0,0), rgba(0,0,0,0.06), rgba(0,0,0,0)) !important;
        margin: 1.2rem 0 !important;
    }

    /* 4. 折叠框 (展开核心摘要) 完美融合 */
    [data-testid="stExpander"] {
        background: rgba(241, 245, 249, 0.5) !important; 
        border: 1px solid rgba(0, 0, 0, 0.03) !important;
        border-radius: 12px !important;
        box-shadow: none !important;
    }
    [data-testid="stExpander"] summary {
        color: #64748B !important;
        font-weight: 600 !important;
        font-size: 14px !important;
    }
    [data-testid="stExpander"] summary:hover {
        color: #334155 !important;
    }

    /* 5. 战术按钮 (阻尼质感) */
    .stButton > button {
        border-radius: 12px !important;
        background: rgba(255, 255, 255, 0.9) !important;
        border: 1px solid rgba(0, 0, 0, 0.05) !important;
        color: #475569 !important;
        font-weight: 600 !important;
        font-size: 14px !important;
        transition: transform 0.3s cubic-bezier(0.34, 1.56, 0.64, 1), box-shadow 0.3s ease !important;
    }
    .stButton > button:hover {
        transform: scale(1.03) !important;
        box-shadow: 0 4px 12px rgba(0, 0, 0, 0.05) !important;
        color: #0F172A !important;
    }
    .stButton > button:active {
        transform: scale(0.96) !important; 
    }
    
    .stButton > button[kind="primary"] {
        background: #64748B !important; 
        color: white !important;
        border: none !important;
    }
    .stButton > button[kind="primary"]:hover {
        background: #475569 !important;
        box-shadow: 0 6px 16px rgba(100, 116, 139, 0.3) !important;
    }

    /* 6. 输入框 (批示区) */
    .stTextInput > div > div > input {
        border-radius: 10px !important;
        background: rgba(255, 255, 255, 0.7) !important;
        border: 1px solid rgba(0, 0, 0, 0.06) !important;
        color: #1E293B !important;
        padding: 0.5rem 1rem !important;
    }
    .stTextInput > div > div > input:focus {
        background: #FFFFFF !important;
        border-color: #94A3B8 !important;
        box-shadow: 0 0 0 2px rgba(148, 163, 184, 0.2) !important;
    }
</style>
""", unsafe_allow_html=True)
SUPABASE_URL = st.secrets["SUPABASE_URL"]
SUPABASE_KEY = st.secrets["SUPABASE_KEY"]
supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)

if "authenticated" not in st.session_state:
    st.session_state.authenticated = False
    st.session_state.current_user = ""
if "page" not in st.session_state:
    st.session_state.page = "main"
if "current_report" not in st.session_state:
    st.session_state.current_report = None
# 🌟 V5.0 新增：用于记录在“专注模式”下，特工当前正在批阅第几张卡片
if "focus_index" not in st.session_state:
    st.session_state.focus_index = 0

if not st.session_state.authenticated:
    st.title("🔒 绝密区域：Agent身份核验")
    st.markdown("---")
    team_members = ['指挥官', '工程师', '梅潮风', '张可可', '丸子', '听风', '瑰夏', '阿拉比卡', '耶加雪非', '蓝山', '曼特宁']
    user_name = st.selectbox("请选择Agent代号：", team_members)
    pwd = st.text_input("请输入访问口令：", type="password")
    if st.button("登录指挥中心", type="primary"):
        auth_response = supabase.table("agents_db").select("*").eq("agent_name", user_name).eq("password", pwd).execute()
        if len(auth_response.data) > 0:
            st.session_state.authenticated = True
            st.session_state.current_user = user_name
            st.success(f"✅ 身份核验通过！欢迎Agent：{user_name}")
            st.rerun()
        else:
            st.error("🚨 警告：授权失败！口令错误或身份不符！")
    st.stop() 

# ==========================================
# 🌸 2. 核心配置与工具函数
# ==========================================
DEEPSEEK_API_KEY = st.secrets["DEEPSEEK_API_KEY"]
ANTHROPIC_API_KEY = st.secrets["ANTHROPIC_API_KEY"]

channel_urls = [
    "https://t.me/s/ejdailyru","https://t.me/s/Ateobreaking", "https://t.me/s/theinsider", "https://t.me/s/moscowtimes_ru",
    "https://t.me/s/economica","https://t.me/s/rybar_africa","https://t.me/s/zakupki_time","https://t.me/s/truestorymedia",
    "https://t.me/s/AoMurmansk","https://t.me/s/moscow_laundry","https://t.me/s/svtvnews","https://t.me/s/notes_veterans",
    "https://t.me/s/militarysummary","https://t.me/s/Tolo_news","https://t.me/s/kremlin_sekret","https://t.me/s/dva_majors",
    "https://t.me/s/caucasar","https://t.me/s/rybar","https://t.me/s/olen_nn","https://t.me/s/russicaRU",
    "https://t.me/s/topwar_official","https://t.me/s/RusskajaIdea","https://t.me/s/riakatysha","https://t.me/s/rybar_latam",
    "https://t.me/s/zhivoff","https://t.me/s/anserenko","https://t.me/s/wolframiumZ","https://t.me/s/vatnoeboloto","https://t.me/s/romanromachev",
    "https://t.me/s/thehegemonist","https://t.me/s/budni_manipulyatora","https://t.me/s/ManoiloToday","https://t.me/s/rtechnocom",
    "https://t.me/s/darpaandcia","https://t.me/s/istories_media","https://t.me/s/mediazona_exclusive","https://t.me/s/Russian_OSINT",
    "https://t.me/s/alter_academy","https://t.me/s/rybar_mena","https://t.me/s/rybar_pacific","https://t.me/s/mosnews","https://t.me/s/brieflyru"
    "https://t.me/s/FarsNewsInt","https://t.me/s/borisenkoD","https://t.me/s/Shiryaev_and_Shiryaev","https://t.me/s/rusbri",
    "https://t.me/s/MedvedevVesti","https://t.me/s/rybar_america","https://t.me/s/russ_orientalist","https://t.me/s/voinasordoy",]  
VIP_CHANNELS = ["anserenko", "kremlin_sekret","https://t.me/s/rybar_america","rybar","Russian_OSINT","rybar_mena","rybar_pacific","topwar_official"] 

def load_bookmarks():
    try:
        res = supabase.table("bookmarks_db").select("*").execute()
        return {row['channel_name']: row['last_read_id'] for row in res.data}
    except: return {}

def save_bookmarks(bookmarks):
    try:
        data = [{"channel_name": k, "last_read_id": v} for k, v in bookmarks.items()]
        if data: supabase.table("bookmarks_db").upsert(data).execute()
    except: pass

def generate_word_doc(title, content):
    doc = Document()
    doc.add_heading(f"深度研判专报：{title}", 0)
    doc.add_paragraph(content)
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

# ==========================================
# 🌸 3. 全局左侧战术侧边栏 (多页面导航)
# ==========================================
with st.sidebar:
    st.title("⚙️ 战术控制台")
    st.success(f"🟢 在线Agent: **{st.session_state.current_user}**")
    
    st.markdown("---")
    st.subheader("🧭 导航中心")
    if st.button("🏠 主力指挥大厅", use_container_width=True, type="primary" if st.session_state.page == "main" else "secondary"):
        st.session_state.page = "main"
        st.rerun()
    if st.button("🗂️ 情报归档库", use_container_width=True, type="primary" if st.session_state.page == "archives" else "secondary"):
        st.session_state.page = "archives"
        st.rerun()
    if st.button("👁️ 深渊挖掘档案室", use_container_width=True, type="primary" if st.session_state.page == "deep_dive_list" else "secondary"):
        st.session_state.page = "deep_dive_list"
        st.rerun()
    if st.button("🗓️ 战略内参简报", use_container_width=True):
        st.session_state.page = "briefings"
        st.rerun()
    st.markdown("---")
    st.caption("🌸 花魁 OSINT v4.1 | 协同归档版")

# ==========================================
# 🌸 4. 页面 1：主力指挥大厅 (抓取与展示)
# ==========================================
if st.session_state.page == "main":
    with st.sidebar:
        st.markdown("---")
        run_btn = st.button("🚀 启动常态挖掘 (DeepSeek)", use_container_width=True, type="primary")
        st.subheader("🎯 实时大厅筛选器")
        filter_category = st.selectbox("领域锁定：", ["全部领域", "China Nexus", "Espionage", "Kremlin Core", "RU Local Event", "Global Macro"])
        filter_score = st.slider("最低威胁分阀值：", 0, 100, 0)

    if run_btn:
        with st.spinner('调用 DeepSeek 引擎执行广度侦察（含时间轴穿透）...'):
            try:
                bookmarks = load_bookmarks()
                raw_intelligence = ""
                new_msg_count = 0
                headers = {'User-Agent': 'Mozilla/5.0'}
                
                for url in channel_urls:
                    try:
                        channel_name = url.split('/s/')[-1]
                        last_read_id = bookmarks.get(channel_name, 0)
                        highest_id = last_read_id
                        
                        channel_new_text = ""
                        current_url = url
                        pages_fetched = 0
                        
                        while pages_fetched < 5:
                            response = requests.get(current_url, headers=headers)
                            soup = BeautifulSoup(response.text, 'html.parser')
                            message_blocks = soup.find_all('div', class_='tgme_widget_message')
                            
                            if not message_blocks: break
                            
                            if last_read_id == 0 and pages_fetched == 0: 
                                message_blocks = message_blocks[-20:] 
                                
                            page_has_new_msg = False
                            oldest_msg_id_on_page = 999999999
                            page_text = ""
                            
                            for block in message_blocks:
                                post_id_str = block.get('data-post')
                                text_div = block.find('div', class_='tgme_widget_message_text')
                                time_tag = block.find('time')
                                msg_time = time_tag.get('datetime', '')[:16].replace('T', ' ') if time_tag else "未知时间"
                                
                                if post_id_str and text_div:
                                    msg_id = int(post_id_str.split('/')[-1])
                                    if msg_id < oldest_msg_id_on_page: 
                                        oldest_msg_id_on_page = msg_id
                                    
                                    if msg_id > last_read_id:
                                        page_text += f"[发帖时间: {msg_time}] " + text_div.text + "\n"
                                        new_msg_count += 1
                                        page_has_new_msg = True
                                        if msg_id > highest_id: highest_id = msg_id
                            
                            channel_new_text = page_text + channel_new_text
                            
                            if oldest_msg_id_on_page > last_read_id and page_has_new_msg and last_read_id != 0:
                                current_url = f"{url}?before={oldest_msg_id_on_page}"
                                pages_fetched += 1
                            else:
                                break
                                
                        if channel_new_text != "":
                            is_vip = "【🔴 VIP 必须提炼】" if channel_name in VIP_CHANNELS else ""
                            raw_intelligence += f"\n\n--- 来源：{channel_name} {is_vip} ---\n" + channel_new_text
                            bookmarks[channel_name] = highest_id
                    except: pass 
                save_bookmarks(bookmarks)
                
                if new_msg_count == 0:
                    st.sidebar.success("暂无更新。")
                else:
                    client = OpenAI(api_key=DEEPSEEK_API_KEY, base_url="https://api.deepseek.com")
                    system_prompt = """
                    你是一位顶级的地缘政治与 OSINT 分析官。
                    请分析原始消息，浓缩成独立情报，必须彻底翻译为简体中文！
                    
                    ⚠️ 极其重要指令 1：原始文本中带有 [发帖时间: ...]。如果有多个来源讲述同一件事，请提取出其中最早的那个时间，格式为 YYYY-MM-DD HH:MM。
                    ⚠️ 极其重要指令 2：如果该条信息的来源带有 "【🔴 VIP 必须提炼】" 的标记，请在输出的 "summary" 字段最后，追加 "【💎 VIP 原文全译】：" 及完整的中文翻译。
                    ⚠️ 极其重要指令 3（防崩溃最高纪律）：你输出的必须是严格合法的 JSON！JSON 字符串内部的换行必须严格使用 "\\n" 代替，绝对不能直接物理换行！双引号必须用 "\\" 转义！
                    ⚠️ 极其重要指令 4：请务必“宁滥勿缺”！只要包含具体事件、观点、动向，即便你认为价值不高，也要提取出来（可以打低分），绝不能随意丢弃原始信息！

                    【🎯 核心战术打分量表 (score: 0-100)】：
                    - 90-100分 (极高危/战略级)：将改变地缘格局、重大高层清洗/人事突变、涉华重大负面/核心利益链异动、核潜艇/战略武器调动。
                    - 70-89分 (高价值线索)：中等规模突发冲突、关键供应链/能源网异动。
                    - 40-69分 (一般情报)：常规战况播报、例行外交辞令、宏观经济数据的一般波动。
                    - 0-39分 (信息噪点)：无意义的政治宣传、未经证实的边缘八卦、日常琐事。（请尽量将此类信息剔除，不要输出）。
                    
                    【输出合法 JSON】：
                    {
                        "reports": [
                            {"title": "中文标题", "summary": "中文概述及VIP全文（注：换行必须用 \\n，注意转义双引号）", "category": "China Nexus 等代号", "score": 85, "source": "频道", "publish_time": "最早发布时间(YYYY-MM-DD HH:MM)"}
                        ]
                    }
                    """
                    ai_response = client.chat.completions.create(
                        model="deepseek-chat", messages=[{"role": "system", "content": system_prompt}, {"role": "user", "content": raw_intelligence}],
                        response_format={"type": "json_object"}, max_tokens=8000
                    )
                    reports = json.loads(ai_response.choices[0].message.content).get("reports", [])
                    for rep in reports:
                        supabase.table("intelligence_db").insert({
                            "title": rep.get("title", "无标题"), "summary": rep.get("summary", "无内容"),
                            "category": rep.get("category", "Global Macro"), "score": rep.get("score", 0), "source": rep.get("source", "未知"),
                            "publish_time": rep.get("publish_time", "未知时间")
                        }).execute()
                    st.sidebar.success(f"✅ 截获 {len(reports)} 条中文情报！")
            except Exception as e: st.error(f"故障：{e}")

    st.title("🌸 OSINT 指挥大厅 (实时截获)")
    
    # 🌟 V5.0 核心大招：UI 模式无缝切换开关
    ui_mode = st.radio("👁️ 战术视觉模式切换：", ["初始模式 (经典列表)", "信息瀑布模式 (全局视野)", "专注模式 (沉浸审批)"], horizontal=True)
    st.markdown("---")
    
    try:
        db_response = supabase.table("intelligence_db").select("*").order("id", desc=True).execute()
        db_cards = db_response.data
    except: db_cards = []

    if len(db_cards) > 0:
        filtered_cards = [c for c in db_cards if (filter_category == "全部领域" or c.get('category') == filter_category) and c.get('score', 0) >= filter_score]
        
        if len(filtered_cards) == 0:
            st.info("💡 当前筛选条件下无匹配情报。")
        else:
            # ==========================================
            # 模式 1：初始模式 (经典单列列表，原汁原味)
            # ==========================================
            if ui_mode == "初始模式 (经典列表)":
                for card in filtered_cards:
                    score = card.get('score', 0)
                    border_color = "🔴" if score >= 80 else "🟡" if score >= 60 else "🔵"
                    with st.container(border=True):
                        st.markdown(f"### {border_color} [{score}分] {card.get('category')} | {card.get('title')}")
                        st.caption(f"📡 来源：{card.get('source')} | 🕰️ 真实发布时间：**{card.get('publish_time', '未获取')}**")
                        st.write(card.get('summary'))
                        
                        comments_res = supabase.table("comments_db").select("*").eq("report_id", card['id']).order("created_at").execute()
                        if len(comments_res.data) > 0:
                            st.markdown("---")
                            for c in comments_res.data:
                                st.markdown(f"**🕵️ {c['agent_name']}** : {c['content']}")
                        st.markdown("---")
                        
                        c1, c2, c3, c4 = st.columns([2, 1, 1, 1])
                        with c1:
                            comment_text = st.text_input("📝 批示...", key=f"in_{card['id']}", label_visibility="collapsed")
                        with c2:
                            if st.button("💬 提交批示", key=f"btn_c_{card['id']}", use_container_width=True) and comment_text:
                                supabase.table("comments_db").insert({"report_id": card['id'], "agent_name": st.session_state.current_user, "content": comment_text}).execute()
                                st.rerun()
                        with c3:
                            if st.button("🔍 深度挖掘", key=f"btn_d_{card['id']}", use_container_width=True, type="secondary"):
                                st.session_state.current_report = card
                                st.session_state.page = "deep_dive" 
                                st.rerun()
                        with c4:
                            if st.button("⭐ 归档入库", key=f"btn_arc_{card['id']}", use_container_width=True):
                                try:
                                    supabase.table("archives_db").insert({"report_id": card['id'], "agent_name": st.session_state.current_user}).execute()
                                    st.toast("✅ 成功存入您的私人归档库！")
                                except:
                                    st.toast("⚠️ 这条情报您之前已经归档过了！")
            
            # ==========================================
            # 模式 2：信息瀑布模式 (三列高低错落排布)
            # ==========================================
            elif ui_mode == "信息瀑布模式 (全局视野)":
                cols = st.columns(3) # 创建 3 列
                for i, card in enumerate(filtered_cards):
                    col = cols[i % 3] # 轮流把卡片塞进 3 列里
                    with col.container(border=True):
                        score = card.get('score', 0)
                        border_color = "🔴" if score >= 80 else "🟡" if score >= 60 else "🔵"
                        
                        st.markdown(f"#### {border_color} [{score}分] {card.get('category')}")
                        st.markdown(f"**{card.get('title')}**")
                        st.caption(f"📡 {card.get('source')} | 🕰️ {card.get('publish_time', '未知')}")
                        
                        with st.expander("展开核心摘要与原文"):
                            st.write(card.get('summary'))
                        
                        c_left, c_right = st.columns(2)
                        with c_left:
                            if st.button("🔍 研判", key=f"wf_d_{card['id']}", use_container_width=True):
                                st.session_state.current_report = card
                                st.session_state.page = "deep_dive" 
                                st.rerun()
                        with c_right:
                            if st.button("⭐ 归档", key=f"wf_arc_{card['id']}", use_container_width=True):
                                try:
                                    supabase.table("archives_db").insert({"report_id": card['id'], "agent_name": st.session_state.current_user}).execute()
                                    st.toast("✅ 存入私人归档库！")
                                except: st.toast("⚠️ 已归档过！")

            # ==========================================
            # 模式 3：专注模式 (单卡片沉浸审批)
            # ==========================================
            elif ui_mode == "专注模式 (沉浸审批)":
                if st.session_state.focus_index >= len(filtered_cards):
                    st.session_state.focus_index = 0
                
                card = filtered_cards[st.session_state.focus_index]
                score = card.get('score', 0)
                border_color = "🔴" if score >= 80 else "🟡" if score >= 60 else "🔵"
                
                spacer1, center_col, spacer2 = st.columns([1, 2, 1]) # 使用空列挤压，让卡片居中独占
                
                with center_col.container(border=True):
                    st.markdown(f"## {border_color} [{score}分] {card.get('category')}")
                    st.markdown(f"### {card.get('title')}")
                    st.caption(f"📡 来源：{card.get('source')} | 🕰️ 真实发布时间：**{card.get('publish_time', '未获取')}**")
                    st.markdown("---")
                    st.write(card.get('summary'))
                    st.markdown("---")
                    
                    comments_res = supabase.table("comments_db").select("*").eq("report_id", card['id']).order("created_at").execute()
                    if len(comments_res.data) > 0:
                        st.caption("💬 **战术批示区：**")
                        for c in comments_res.data:
                            st.markdown(f"**🕵️ {c['agent_name']}** : {c['content']}")
                        st.markdown("---")
                    
                    comment_text = st.text_input("📝 添加批示...", key=f"foc_in_{card['id']}", label_visibility="collapsed")
                    if st.button("💬 提交批示并留在本页", key=f"foc_btn_c_{card['id']}", use_container_width=True) and comment_text:
                        supabase.table("comments_db").insert({"report_id": card['id'], "agent_name": st.session_state.current_user, "content": comment_text}).execute()
                        st.rerun()
                    
                    st.markdown("<br>", unsafe_allow_html=True)
                    
                    t1, t2 = st.columns(2)
                    with t1:
                        if st.button("❌ 忽略 / 划走", key=f"foc_pass_{card['id']}", use_container_width=True, type="secondary"):
                            st.session_state.focus_index += 1
                            st.rerun()
                    with t2:
                        if st.button("⭐ 归档入库", key=f"foc_arc_{card['id']}", use_container_width=True, type="primary"):
                            try:
                                supabase.table("archives_db").insert({"report_id": card['id'], "agent_name": st.session_state.current_user}).execute()
                                st.toast("✅ 归档成功！")
                            except: pass
                            st.session_state.focus_index += 1
                            st.rerun()
                            
                    if st.button("🔍 呼叫 Claude 启动深渊研判", key=f"foc_d_{card['id']}", use_container_width=True):
                        st.session_state.current_report = card
                        st.session_state.page = "deep_dive" 
                        st.rerun()
                        
                st.markdown(f"<p style='text-align: center; color: gray;'>👉 审批进度：( {st.session_state.focus_index + 1} / {len(filtered_cards)} )</p>", unsafe_allow_html=True)
    else:
        st.info("👈 报告长官，数据库目前为空。请启动侦察！")

# ==========================================
# 🌸 5. 页面 2：个人与团队情报归档库
# ==========================================
elif st.session_state.page == "archives":
    st.title("🗂️ 战略归档库")
    
    # 顶部选择器：看自己的，还是看全队的？
    view_mode = st.radio("视角切换：", ["👨‍💻 我的私人珍藏", "🌐 团队全员归档公开板"], horizontal=True)
    st.markdown("---")
    
    try:
        # 获取归档关系
        if view_mode == "👨‍💻 我的私人珍藏":
            arch_res = supabase.table("archives_db").select("*").eq("agent_name", st.session_state.current_user).execute()
        else:
            arch_res = supabase.table("archives_db").select("*").execute()
            
        archived_items = arch_res.data
        if len(archived_items) == 0:
            st.info("📦 归档库空空如也，快去指挥大厅淘宝吧！")
        else:
            # 提取所有被归档的 report_id
            report_ids = list(set([item['report_id'] for item in archived_items]))
            
            # 获取情报卡片本体
            db_res = supabase.table("intelligence_db").select("*").in_("id", report_ids).order("id", desc=True).execute()
            
            for card in db_res.data:
                # 找出是谁归档了这张卡（用于全员视角展示）
                archived_by = [a['agent_name'] for a in archived_items if a['report_id'] == card['id']]
                
                with st.container(border=True):
                    st.markdown(f"### [{card.get('score',0)}分] {card.get('title')}")
                    st.caption(f"📡 来源：{card.get('source')} | 🕰️ 真实发布：{card.get('publish_time', '未知')} | 🌟 归档者：`{', '.join(set(archived_by))}`")
                    st.write(card.get('summary'))
                    
                    if st.button("🔍 再次唤醒深度挖掘", key=f"arc_d_{card['id']}"):
                        st.session_state.current_report = card
                        st.session_state.page = "deep_dive"
                        st.rerun()

    except Exception as e:
        st.error(f"读取归档失败: {e}")

# ==========================================
# 🌸 6. 页面 3：深渊挖掘档案室 (已挖掘汇总)
# ==========================================
elif st.session_state.page == "deep_dive_list":
    st.title("👁️ 深渊档案室 (知识资产沉淀)")
    st.caption("这里存放着团队花费重金让大模型运算过的所有《深度研判专报》，全员均可免费查阅。")
    st.markdown("---")
    
    try:
        deep_res = supabase.table("deep_dives_db").select("*").order("created_at", desc=True).execute()
        dives = deep_res.data
        
        if len(dives) == 0:
            st.info("尚无深渊档案。去大厅点击【深度挖掘】生成您的第一份报告！")
        else:
            # 批量获取关联的情报标题
            report_ids = [d['report_id'] for d in dives]
            intel_res = supabase.table("intelligence_db").select("id, title, summary, source").in_("id", report_ids).execute()
            intel_dict = {item['id']: item for item in intel_res.data}
            
            for d in dives:
                origin_card = intel_dict.get(d['report_id'], {})
                with st.expander(f"📁 绝密研判专报：{origin_card.get('title', '未知情报')} (发起特工: {d['agent_name']})"):
                    st.caption(f"提取时间：{d['created_at'][:16].replace('T', ' ')}")
                    st.markdown(d['content'])
                    
                    # 也可以在这里直接下载 Word
                    docx_data = generate_word_doc(origin_card.get('title', ''), d['content'])
                    st.download_button(
                        label="📥 导出此专报 (Word)",
                        data=docx_data,
                        file_name=f"HUMINT_Archive_{d['report_id']}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=f"dl_{d['id']}"
                    )
    except Exception as e:
        st.error(f"读取深渊档案库失败: {e}")

# ==========================================
# 🌸 7. 独立审讯室：Claude 深度研判行动执行
# ==========================================

elif st.session_state.page == "briefings":
    # ==========================================
    # 🌸 V6.1 独立页面：战略内参简报室
    # ==========================================
    st.title("🗓️ 战略内参简报室")
    if st.button("⬅️ 返回战略情报大厅", type="primary"):
        st.session_state.page = "main"
        st.rerun()
        
    st.markdown("---")
    
    try:
        # 倒序拉取所有由 Claude 生成的简报（最多显示最近10期）
        res = supabase.table("briefings_db").select("*").order("created_at", desc=True).limit(10).execute()
        briefings = res.data
        
        if not briefings:
            st.info("📭 目前档案库中还没有生成任何简报。请等待每日早 8 点系统自动生成，或前往 GitHub Actions 手动发射无人机触发。")
        else:
            for b in briefings:
                period_icon = "🌅" if b['period'] == 'daily' else "🗓️"
                # 提取简报的第一行作为手风琴折叠框的标题
                title_line = b['content'].split('\n')[0].replace('# ', '')
                with st.expander(f"{period_icon} {title_line}", expanded=False):
                    st.markdown(b['content'])
    except Exception as e:
        st.error(f"调取简报档案库失败，数据库链接异常：{e}")

elif st.session_state.page == "deep_dive":
    # ==========================================
    # 🌸 7. 独立战术研判室：Claude Agent (联网搜索 + 军工级模板)
    # ==========================================
    card = st.session_state.current_report
    st.title("👁️ 独立战术研判室 (Agent 破壁版)")
    if st.button("⬅️ 返回战略情报大厅", type="primary"):
        st.session_state.page = "main"
        st.rerun()
        
    st.markdown("---")
    st.markdown(f"#### 【原始截获情报】\n**{card['title']}**\n> {card['summary']}")
    
    deep_res = supabase.table("deep_dives_db").select("*").eq("report_id", card['id']).execute()
    
    if len(deep_res.data) > 0:
        st.success(f"💾 历史档案调取成功！本报告由Agent **{deep_res.data[0]['agent_name']}** 耗费 AI 算力挖掘，现为您免费呈现。")
        final_content = deep_res.data[0]['content']
        st.markdown(final_content)
        
    else:
        with st.spinner("🧠 正在启动 Agent 智能体：执行外网交叉检索与 Claude 深层推演..."):
            try:
                # 🌟 第一重破壁：使用 DuckDuckGo 搜索真实世界的最新进展
                search_results_text = "【外网检索结果】：\n"
                try:
                    search_query = f"{card['title']} {card['category']} news"
                    with DDGS() as ddgs:
                        # 抓取排名前 3 的真实外网数据
                        results = [r for r in ddgs.text(search_query, max_results=3)]
                        if results:
                            for i, r in enumerate(results):
                                search_results_text += f"{i+1}. {r.get('title')} - {r.get('body')}\n"
                        else:
                            search_results_text += "未能在公开网络检索到相关增量信息（极高保密级别或为虚假信息）。\n"
                except Exception as search_e:
                    search_results_text += f"雷达受到干扰，检索失败 ({search_e})。仅依赖基础逻辑推演。\n"

                # 🌟 第二重破壁：内参级智库专报模板 (彻底杀死 AI 味，锁定战略干货)
                # 🌟 第三重破壁：内参级思维链 (CoT) 研判大脑，彻底消灭 AI 味
                claude_system_prompt = """
                你是一名服务于中国最高决策层的“首席地缘战略与情报分析官”。你极其冷酷、理性、毒辣，能一眼看穿政治宣传背后的真实利益交换。
                ⚠️ 【绝对禁令】：彻底消除一切 AI 痕迹。严禁使用“总而言之”、“值得注意的是”、“作为一个AI”、“综上所述”等公文废话。你的文字必须像剔骨刀一样，通篇只有高密度的干货与事实推演。全文严格控制在 2000 字以内。

                【工作流指令 - 强制思维链 (Chain of Thought)】：
                在撰写正式报告前，你必须首先在一个 <thinking> ... </thinking> 的 XML 标签内进行深度的内部逻辑推演。
                你的推演必须包含：
                1. 交叉比对：初始情报与 DuckDuckGo 外网搜索结果是否有矛盾？谁在撒谎？谁在放风？
                2. 动机反推：这个消息在这个时间点被放出来，最大受益者是谁？这是否是一场认知战（舆论战）？
                3. 底线推演：事件的最坏发展路径是什么？

                思考完毕后，离开 <thinking> 标签，严格按照以下《绝密战略内参模板》输出正式报告：

                # [生成一个极具穿透力的标题，格式如：“俄罗斯总统办公室对当前伊朗局势发展的分析判断”]

                ## 核心结论 (BLUF)
                [用极其冷峻、不带感情色彩的 1-2 句话，直接揭示该事件的本质、目前的真实走向，以及对核心地缘格局的最终影响。让决策者 3 秒内看懂核心危机。]

                ## 一、 事实复盘与情报证伪
                [结合初始线索与外网搜索结果，还原事件的真实时间线。必须明确指出：外网搜索到的证据是“印证”、“证伪”还是“补充”了初始情报？剥离掉各方的新闻宣传，只留下客观发生的动作、资金、人事或军力调动。]

                ## 二、 幕后暗战与各方底牌
                [基于你在思考阶段的推演，极其深入地剖析事件背后至少两到三个核心利益方的真实考量。不要写他们“声称”什么，只写他们“害怕”什么、“想要”什么、他们下一步可能的“隐秘动作”是什么。必须有逻辑支撑，杜绝主观臆测。]

                ## 三、 涉华研判与政策级对策
                [抛弃空泛的“外交呼吁”。从中国国家利益（能源安全、地缘屏障、金融避险、供应链等）出发，评估其造成的实质性威胁或机遇。并给出至少 3 条极其具体、可落地的国家级应对建议（例如：建议某部委采取何种经贸对冲、建议军方关注某个特定海峡、建议情报网加强对某个特定寡头的监控）。]
                """
                
                claude_user_message = f"""
                【初始拦截线索】：
                - 标题：{card['title']}
                - 摘要：{card['summary']}
                - 来源：{card['source']}
                
                【外网交叉验证搜索结果】：
                {search_results_text}
                
                请立刻开始执行：先在 <thinking> 中进行逻辑绞杀与沙盘推演，然后输出《绝密战略内参》。
                """
                
                client_claude = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
                response = client_claude.messages.create(
                    model="claude-sonnet-4-6", # 或者最新的 claude-3-5-sonnet-latest
                    max_tokens=4096,
                    system=claude_system_prompt,
                    messages=[{"role": "user", "content": claude_user_message}]
                )
                final_content = response.content[0].text
                
                # 记录挖掘结果
                supabase.table("deep_dives_db").insert({
                    "report_id": card['id'], "agent_name": st.session_state.current_user, "content": final_content
                }).execute()
                
                st.success(f"🔥 Agent 智能体检索并研判完毕！已刻录至深渊档案库。")
                
                # 显示外网检索过程（给特工看一眼雷达抓到了什么）
                with st.expander("📡 查看系统后台外网检索原始数据"):
                    st.text(search_results_text)
                    
                st.markdown(final_content)
            except Exception as e:
                st.error(f"Agent 引擎严重故障：{e}")
                final_content = ""

    if final_content:
        st.markdown("---")
        docx_data = generate_word_doc(card['title'], final_content)
        st.download_button(
            label="📥 导出为 Word 绝密专报 (.docx)", data=docx_data, file_name=f"HUMINT专报_{card['id']}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", type="primary"
        )